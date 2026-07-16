"""Extract paragraphs, headings, tables, and images from DOCX files (summaries + past exam docx)."""
import json
import sys
from pathlib import Path

import docx
from docx.document import Document as DocumentObj
from docx.oxml.ns import qn

from config import EXTRACTED, SUMMARIES_DIR, EXAMS_DIR, slugify


def extract_images(document, img_dir: Path):
    img_dir.mkdir(parents=True, exist_ok=True)
    paths = []
    try:
        rels = document.part.rels
        i = 0
        for rel in rels.values():
            if "image" in rel.reltype:
                i += 1
                try:
                    ext = rel.target_ref.rsplit(".", 1)[-1]
                    data = rel.target_part.blob
                    p = img_dir / f"img_{i:03d}.{ext}"
                    p.write_bytes(data)
                    paths.append(str(p.relative_to(EXTRACTED.parent)))
                except Exception as e:
                    paths.append(f"ERROR: {e}")
    except Exception as e:
        paths.append(f"ERROR listing rels: {e}")
    return paths


def extract_one(docx_path: Path, out_dir: Path, category: str) -> dict:
    file_id = slugify(docx_path.name)
    document = docx.Document(str(docx_path))

    blocks = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        blocks.append({"type": "paragraph", "style": style, "text": text})

    tables = []
    for t_idx, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"tableIndex": t_idx, "rows": rows})

    img_dir = EXTRACTED / "images" / file_id
    images = extract_images(document, img_dir)

    word_count = sum(len(b["text"].split()) for b in blocks)

    result = {
        "id": file_id,
        "sourcePath": str(docx_path),
        "category": category,
        "fileName": docx_path.name,
        "paragraphBlocks": blocks,
        "tables": tables,
        "images": images,
        "wordCount": word_count,
    }

    out_path = out_dir / f"{file_id}.json"
    out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    return {
        "id": file_id,
        "sourcePath": str(docx_path),
        "outputPath": str(out_path),
        "wordCount": word_count,
        "paragraphCount": len(blocks),
        "tableCount": len(tables),
        "imageCount": len(images),
        "status": "extracted",
    }


def main():
    jobs = []
    for f in sorted(SUMMARIES_DIR.glob("*.docx")):
        jobs.append((f, EXTRACTED / "summaries", "summary"))
    for f in sorted(EXAMS_DIR.glob("*.docx")):
        jobs.append((f, EXTRACTED / "exams", "past_exam"))

    manifest_entries = []
    for docx_path, out_dir, category in jobs:
        try:
            entry = extract_one(docx_path, out_dir, category)
            print(f"OK  {category:10s} {docx_path.name} -> {entry['wordCount']} words, {entry['tableCount']} tables, {entry['imageCount']} images")
        except Exception as e:
            entry = {"id": slugify(docx_path.name), "sourcePath": str(docx_path), "status": "error", "error": str(e)}
            print(f"ERR {category:10s} {docx_path.name}: {e}", file=sys.stderr)
        manifest_entries.append(entry)

    manifest_path = EXTRACTED / "docx_manifest.json"
    manifest_path.write_text(json.dumps(manifest_entries, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\nWrote {manifest_path} with {len(manifest_entries)} entries")


if __name__ == "__main__":
    main()
